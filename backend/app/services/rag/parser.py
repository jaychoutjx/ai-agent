"""
文档解析层：把不同格式的文档统一解析成纯文本。

【支持格式】
- PDF（.pdf）       → pypdf
- Word（.docx）     → python-docx
- Markdown（.md）   → 直接读取
- 纯文本（.txt）    → 直接读取
- Excel（.xlsx）    → openpyxl

【面试可讲】为什么用工厂模式？
- 新增格式只需加一个解析函数和注册一行，不改主流程
- 调用方只关心 parse_document(path)，不关心内部如何解析
"""

from collections.abc import Callable
from pathlib import Path

from app.core.logger import logger


def _parse_pdf(path: Path) -> str:
    """解析 PDF。每页文本拼接，页之间用换行分隔。"""
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages: list[str] = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"\n[第 {i + 1} 页]\n{text.strip()}")
    return "\n\n".join(pages)


def _parse_docx(path: Path) -> str:
    """解析 Word 文档（.docx）。提取所有段落 + 表格内容。"""
    from docx import Document

    doc = Document(str(path))

    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _parse_xlsx(path: Path) -> str:
    """解析 Excel。每个工作表用标题分隔，每行用 | 拼接。"""
    from openpyxl import load_workbook

    wb = load_workbook(str(path), data_only=True, read_only=True)
    parts: list[str] = []
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        parts.append(f"\n[Sheet: {sheet_name}]")
        for row in sheet.iter_rows(values_only=True):
            row_text = " | ".join(str(c) for c in row if c is not None)
            if row_text.strip():
                parts.append(row_text)
    wb.close()
    return "\n".join(parts)


def _parse_text(path: Path) -> str:
    """直接读取纯文本/Markdown。"""
    return path.read_text(encoding="utf-8", errors="ignore")


PARSERS: dict[str, Callable[[Path], str]] = {
    ".pdf": _parse_pdf,
    ".docx": _parse_docx,
    ".xlsx": _parse_xlsx,
    ".md": _parse_text,
    ".markdown": _parse_text,
    ".txt": _parse_text,
}


SUPPORTED_EXTENSIONS = set(PARSERS.keys())


def parse_document(path: str | Path) -> str:
    """
    统一解析入口。根据文件扩展名分发到对应解析器。

    Args:
        path: 文件路径

    Returns:
        解析后的纯文本

    Raises:
        ValueError: 不支持的文件格式
        FileNotFoundError: 文件不存在
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {path}")

    ext = path.suffix.lower()
    parser = PARSERS.get(ext)
    if parser is None:
        raise ValueError(
            f"不支持的文件格式: {ext}，支持的格式: {sorted(SUPPORTED_EXTENSIONS)}"
        )

    logger.info(f"解析文档: {path.name} ({ext})")
    text = parser(path)
    logger.info(f"解析完成: {path.name}, 文本长度 {len(text)} 字符")
    return text
